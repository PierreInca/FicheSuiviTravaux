from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn

st.write("--- Début de rempl word.py ---")

def appliquer_style(run, taille=8, gras=False, couleur=RGBColor(0, 0, 0)):
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(taille)
    run.font.bold = gras
    run.font.color.rgb = couleur

def remplacer_champs_dans_paragraphes(paragraphs, champs):
    projet1_formate = False  # Pour ne styliser qu'une fois

    for p in paragraphs:
        texte_original = "".join(run.text for run in p.runs)
        texte_modifie = texte_original

        for key, value in champs.items():
            if value is None:
                value = ""
            texte_modifie = texte_modifie.replace(key, value)

        if texte_modifie != texte_original:
            # Supprimer tous les runs existants
            while p.runs:
                p._element.remove(p.runs[0]._element)

            # Traitement spécial pour "PROJET : {{projet1}}"
            if not projet1_formate and champs.get("{{projet1}}"):
                valeur_projet = champs["{{projet1}}"]
                motif = f"PROJET : {valeur_projet}"
                if motif in texte_modifie:
                    avant, sep, apres = texte_modifie.partition(motif)

                    if avant:
                        run_avant = p.add_run(avant)
                        appliquer_style(run_avant)

                    run_cible = p.add_run(motif)
                    appliquer_style(run_cible, taille=12, gras=True, couleur=RGBColor(31, 78, 121))
                    projet1_formate = True

                    if apres:
                        run_apres = p.add_run(apres)
                        appliquer_style(run_apres)

                    continue  # Passer au paragraphe suivant, c'est traité

            # Si ce n'est pas le cas spécial → un seul run avec style par défaut
            run = p.add_run(texte_modifie)
            appliquer_style(run)

            # Cas spécial pour num_fiche (ex. fond blanc sur blanc)
            if champs.get("{{num_fiche}}") and champs["{{num_fiche}}"] in texte_modifie:
                appliquer_style(run, taille=10, gras=False, couleur=RGBColor(255, 255, 255))

def remplacer_champs_dans_doc(doc: Document, champs: dict):
    remplacer_champs_dans_paragraphes(doc.paragraphs, champs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                remplacer_champs_dans_paragraphes(cell.paragraphs, champs)
